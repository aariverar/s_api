import os
import shutil
from docx import Document
from docx.shared import Inches
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxtpl import DocxTemplate
import random

# Constantes
PATH_RELATIVE_WORD = os.path.join(os.getcwd(), "src", "test", "resources","Plantilla.docx")
TEMPLATE = 'template.png'
WORD_NAME_STANDAR = 'nombre_estandar_archivo_word.docx'
WORD_EXTENSION = '.docx'

def start_up_word(rutaEvidencia):
    insert_template = None
    try:
        
        # Crear una carpeta temporal para guardar el archivo de Word
        carpeta = 'Evidencias'
        os.makedirs(carpeta, exist_ok=True)

        if not os.path.exists(rutaEvidencia):
            doc = Document(PATH_RELATIVE_WORD)
            print("[LOG] Se crea word de plantilla")
        else:
            doc = Document(rutaEvidencia)
            print("[LOG] Abre word creado")
        
        return doc
    except Exception as e:
        print(f"[ERROR] {e}")
    finally:
        if insert_template:
            insert_template.close()

def guardarWord(xword, ruta_evidencia_word):
    try:
        # Renombrar el archivo con el estado final
        xword.save(ruta_evidencia_word)

        print("[LOG] Word Guardado")

    except Exception as e:
        print(f"[ERROR] {e}")

def crear_tabla_inicio(doc, context):
    # Añadir un párrafo al final del documento
    #doc.add_paragraph()
    ROW1=3
    COL1=4
    # Establecer la fuente y el tamaño del texto
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # Añadir una tabla con 3 filas y 4 columnas
    table = doc.add_table(rows=ROW1, cols=COL1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.style = 'Table Grid'
    # Ajustar el ancho de las columnas
    for column in table.columns:
        for cell in column.cells:
            cell.width = doc.sections[0].page_width / 4.6

    # Añadir bordes y alinear verticalmente las celdas
    for i in range(ROW1):
        for o in range(COL1):
            table.cell(i, o).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


    # Rellenar la tabla con los datos
    set_cell_text(table.cell(0,0),"QA Agile: ",bold=True,background_color=colorRGB("GRIS"))
    set_cell_text(table.cell(0,1),"Equipo de automatización")

    set_cell_text(table.cell(0,2),"Fecha Ejecución: ",bold=True,background_color=colorRGB("GRIS"))
    set_cell_text(table.cell(0,3),generar_fecha())

    set_cell_text(table.cell(1,0),"Servicio: ",bold=True,background_color=colorRGB("GRIS"))
    table.cell(1, 1).merge(table.cell(1, 3))
    set_cell_text(table.cell(1,1),context.nameApi)

    if context.ciclo=="-":
        set_cell_text(table.cell(2,0),"Estado: ",bold=True,background_color=colorRGB("GRIS"))
        table.cell(2,1).merge(table.cell(2,3))
        set_cell_text(table.cell(2,1),"")
    else:
        set_cell_text(table.cell(2,0),"Ciclo: ",bold=True,background_color=colorRGB("GRIS"))
        set_cell_text(table.cell(2,1),context.ciclo)

        set_cell_text(table.cell(2,2),"Estado: ",bold=True,background_color=colorRGB("GRIS"))
        set_cell_text(table.cell(2,3),"")

    # Añadir un párrafo al final del documento
    doc.add_paragraph()
    print("[LOG] Se crea tabla Inicio")
   
def set_status_scenario(status,context):
    document = context.xdoc
    table = document.tables[0]
    if context.ciclo=="-":
        set_cell_text(table.cell(2,1),status.upper(),background_color=colorStatus(status.upper()))
    else:
        set_cell_text(table.cell(2,3),status.upper(),background_color=colorStatus(status.upper()))
    guardarWord(document , context.rutaEvidencia)
    file_with_new_name = f"{context.rutaEvidencia.split('.docx')[0]}-{status.upper()}{WORD_EXTENSION}"
    os.rename(context.rutaEvidencia, file_with_new_name)

def crear_tabla_response(doc,iteracion, caso, metodo, endPoint, status_response, value_response,headers,body,response,time_response):
      # Añadir un párrafo al final del documento
    #doc.add_paragraph()
    ROW1=7
    COL1=4
    # Establecer la fuente y el tamaño del texto
    p = doc.add_paragraph()
    run = p.add_run(f"Evidencia de Pruebas {iteracion}:")
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.bold= True

    # Añadir una tabla con 3 filas y 4 columnas
    table = doc.add_table(rows=ROW1, cols=COL1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.style = 'Table Grid'
    # Ajustar el ancho de las columnas
    for column in table.columns:
        for cell in column.cells:
            cell.width = doc.sections[0].page_width / 4.6
    table.autofit = False
    # Añadir bordes y alinear verticalmente las celdas
    for i in range(ROW1):
        for o in range(COL1):
            table.cell(i, o).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


    # Rellenar la tabla con los datos
    set_cell_text(table.cell(0,0),"Caso: ",bold=True,background_color='D3D3D3')
    set_cell_text(table.cell(0,1),caso)

    set_cell_text(table.cell(0,2),"Método: ",bold=True,background_color='D3D3D3')
    set_cell_text(table.cell(0,3),metodo)

    set_cell_text(table.cell(1,0),"EndPoint: ",bold=True,background_color='D3D3D3')
    table.cell(1, 1).merge(table.cell(1, 3))
    set_cell_text(table.cell(1,1),endPoint)

    set_cell_text(table.cell(2,0),"Status Response: ",bold=True,background_color='D3D3D3')
    set_cell_text(table.cell(2,1),status_response)

    set_cell_text(table.cell(2,2),"Value Status Response: ",bold=True,background_color='D3D3D3')
    set_cell_text(table.cell(2,3),value_response )

    set_cell_text(table.cell(3,0),"Headers: ",bold=True,background_color='D3D3D3')
    table.cell(3, 1).merge(table.cell(3, 3))
    set_cell_text(table.cell(3,1),headers,alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)

    set_cell_text(table.cell(4,0),"Body: ",bold=True,background_color='D3D3D3')
    table.cell(4, 1).merge(table.cell(4, 3))
    set_cell_text(table.cell(4,1),body,alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)

    set_cell_text(table.cell(5,0),"Response: ",bold=True,background_color='D3D3D3')
    table.cell(5, 1).merge(table.cell(5, 3))
    set_cell_text(table.cell(5,1),response,alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)

    set_cell_text(table.cell(6,0),"Tiempo de Respuesta: ",bold=True,background_color='D3D3D3')
    table.cell(6, 1).merge(table.cell(6, 3))
    set_cell_text(table.cell(6,1),time_response,alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
    # Añadir un párrafo al final del documento
    doc.add_paragraph()
    print("[LOG] Se crea tabla response")

def crear_tabla_validacion(doc,iteracion,status, expected_code, actual_code = "", title=""):
      # Añadir un párrafo al final del documento
    #doc.add_paragraph()
    if actual_code=="":
        ROW1=2
    else:
        ROW1=3
    COL1=2
    # Establecer la fuente y el tamaño del texto
    p = doc.add_paragraph()
    if title=="":
        run = p.add_run(f"Validaciones {iteracion}:")
    else:
        run = p.add_run(f"{title} - {iteracion}")
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.bold= True

    # Añadir una tabla con 3 filas y 4 columnas
    table = doc.add_table(rows=ROW1, cols=COL1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.style = 'Table Grid'
    # Ajustar el ancho de las columnas
    # column_widths = [2.0, 2.5, 2.0, 2.5]  # Ancho de las columnas en pulgadas
    # for i, width in enumerate(column_widths):
    #     for cell in table.columns[i].cells:
    #         cell.width = width * 914400
    table.autofit = False
    for column in table.columns:
        for cell in column.cells:
            cell.width = doc.sections[0].page_width / 4.6

    # Añadir bordes y alinear verticalmente las celdas
    for i in range(ROW1):
        for o in range(COL1):
            table.cell(i, o).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


    # Rellenar la tabla con los datos
    if actual_code=="":
        set_cell_text(table.cell(0,0),"Resultado Esperado: ",bold=True,background_color=colorRGB("GRIS"))
        set_cell_text(table.cell(0,1),expected_code)
        set_cell_text(table.cell(1,0),"Validación: ",bold=True,background_color=colorRGB("GRIS"))
        set_cell_text(table.cell(1,1),status,background_color=colorStatus(status))
    else:
        set_cell_text(table.cell(0,0),"Resultado Esperado: ",bold=True,background_color=colorRGB("GRIS"))
        set_cell_text(table.cell(0,1),expected_code)

        set_cell_text(table.cell(1,0),"Resultado Obtenido: ",bold=True,background_color=colorRGB("GRIS"))
        set_cell_text(table.cell(1,1),actual_code)

        set_cell_text(table.cell(2,0),"Validación: ",bold=True,background_color=colorRGB("GRIS"))
        set_cell_text(table.cell(2,1),status,background_color=colorStatus(status))

    doc.add_paragraph()
    print("[LOG] Se crea tabla Validación")

def crear_tabla_validacion_header(doc, expected_headers, actual_headers, status):
 
    ROW1=3
    COL1=4
 
    p = doc.add_paragraph()
    run = p.add_run(f"Validaciones Headers_Response:")
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.bold= True
    table = doc.add_table(rows=ROW1, cols=COL1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.style = 'Table Grid'
   
    table.autofit = False
    for column in table.columns:
        for cell in column.cells:
            cell.width = doc.sections[0].page_width / 4.6
 
    # Añadir bordes y alinear verticalmente las celdas
    for i in range(ROW1):
        for o in range(COL1):
            table.cell(i, o).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
 
 
    # Rellenar la tabla con los datos
    set_cell_text(table.cell(0,0),"Resultado Esperado: ",bold=True,background_color='D3D3D3')
    table.cell(0, 1).merge(table.cell(0, 3))
    set_cell_text(table.cell(0,1),expected_headers,alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
 
    set_cell_text(table.cell(1,0),"Status Real: ",bold=True,background_color='D3D3D3')
    table.cell(1, 1).merge(table.cell(1, 3))
    set_cell_text(table.cell(1,1),actual_headers,alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
 
    set_cell_text(table.cell(2,0),"Validacion: ",bold=True,background_color='D3D3D3')
    table.cell(2, 1).merge(table.cell(2, 3))
    set_cell_text(table.cell(2,1),status,background_color=colorStatus(status))
 
       
    doc.add_paragraph()
    print("[LOG] Se crea tabla Validacion")

def set_cell_text(cell, text, bold=False, underline=False, background_color=None,alignment =WD_ALIGN_PARAGRAPH.CENTER,tipoLetra='Calibri',tamañoLetra=Pt(10)):
        run = cell.paragraphs[0].add_run(text)
        run.bold = bold
        run.underline = underline
        run.font.name = tipoLetra
        run.font.size = tamañoLetra
        cell.paragraphs[0].paragraph_format.alignment = alignment
        if background_color:
            cell_fill = OxmlElement('w:shd')
            cell_fill.set(qn('w:fill'),background_color)
            cell._element.get_or_add_tcPr().append(cell_fill)
        

def colorRGB(color):
    if color == "GRIS":
        return 'D3D3D3'
    else:
        return RGBColor(255,255,255)
    
def colorStatus(status):
    if status == "PASSED" or status == "FOUND":
        return '68FF00'
    else:
        return 'FF0000' 

def generar_secuencia():
    # Obtener la fecha y hora actual en el formato "dd-MM-yyyy_hh-mm-ss"
    now = datetime.datetime.now()
    formatted_date = now.strftime("%d-%m-%Y_%H-%M-%S")
    return formatted_date


def generar_fecha():
    # Obtener la fecha actual en el formato "dd/MM/yyyy"
    now = datetime.datetime.now()
    formatted_date = now.strftime("%d-%m-%Y")
    return formatted_date

def generar_hora():
    # Obtener la hora actual en el formato "hh:mm:ss"
    now = datetime.datetime.now()
    formatted_time = now.strftime("%H:%M:%S")
    return formatted_time
